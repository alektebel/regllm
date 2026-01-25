#!/usr/bin/env python3
"""
Attachment Handler for Document Processing

Handles file uploads and text extraction from various document formats.
"""

import re
from pathlib import Path
from typing import Dict, Optional, List
from dataclasses import dataclass
import logging

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Represents a processed document."""
    filename: str
    text: str
    metadata: Dict
    chunks: List[str]
    format: str


class AttachmentHandler:
    """Handles file uploads and document processing."""

    SUPPORTED_FORMATS = {'.pdf', '.txt', '.docx', '.md', '.doc'}
    MAX_FILE_SIZE_MB = 50
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 200

    def __init__(self, upload_dir: Optional[Path] = None):
        self.upload_dir = upload_dir or Path("data/uploads")
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported."""
        suffix = Path(file_path).suffix.lower()
        return suffix in self.SUPPORTED_FORMATS

    def get_file_size_mb(self, file_path: str) -> float:
        """Get file size in megabytes."""
        return Path(file_path).stat().st_size / (1024 * 1024)

    def process_upload(
        self,
        file_path: str,
        session_id: Optional[str] = None,
    ) -> ProcessedDocument:
        """
        Process an uploaded file and extract text.

        Args:
            file_path: Path to the uploaded file
            session_id: Optional session identifier

        Returns:
            ProcessedDocument with extracted text and metadata
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {suffix}. Supported: {self.SUPPORTED_FORMATS}")

        file_size = self.get_file_size_mb(file_path)
        if file_size > self.MAX_FILE_SIZE_MB:
            raise ValueError(f"File too large: {file_size:.1f}MB (max: {self.MAX_FILE_SIZE_MB}MB)")

        # Extract text based on format
        if suffix == '.pdf':
            text = self._extract_pdf(path)
        elif suffix == '.docx':
            text = self._extract_docx(path)
        elif suffix == '.doc':
            text = self._extract_doc(path)
        elif suffix in {'.txt', '.md'}:
            text = self._extract_text(path)
        else:
            raise ValueError(f"No extractor for format: {suffix}")

        # Clean text
        text = self._clean_text(text)

        # Create chunks
        chunks = self._chunk_text(text)

        # Build metadata
        metadata = {
            "filename": path.name,
            "format": suffix,
            "size_mb": file_size,
            "char_count": len(text),
            "chunk_count": len(chunks),
            "session_id": session_id,
        }

        return ProcessedDocument(
            filename=path.name,
            text=text,
            metadata=metadata,
            chunks=chunks,
            format=suffix,
        )

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import fitz  # PyMuPDF

            text_parts = []
            with fitz.open(str(path)) as doc:
                for page in doc:
                    text_parts.append(page.get_text())

            return "\n".join(text_parts)

        except ImportError:
            # Fallback to pdfplumber
            try:
                import pdfplumber

                text_parts = []
                with pdfplumber.open(str(path)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)

                return "\n".join(text_parts)

            except ImportError:
                raise ImportError(
                    "PDF processing requires PyMuPDF or pdfplumber. "
                    "Install with: pip install PyMuPDF or pip install pdfplumber"
                )

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document

            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)

        except ImportError:
            raise ImportError(
                "DOCX processing requires python-docx. "
                "Install with: pip install python-docx"
            )

    def _extract_doc(self, path: Path) -> str:
        """Extract text from old DOC format."""
        try:
            import textract

            text = textract.process(str(path)).decode('utf-8')
            return text

        except ImportError:
            raise ImportError(
                "DOC processing requires textract. "
                "Install with: pip install textract"
            )

    def _extract_text(self, path: Path) -> str:
        """Extract text from plain text or markdown file."""
        return path.read_text(encoding='utf-8')

    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Fix common PDF extraction issues
        text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)  # Fix hyphenation

        # Remove null characters
        text = text.replace('\x00', '')

        return text.strip()

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if len(text) <= self.CHUNK_SIZE:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.CHUNK_SIZE

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end within the overlap zone
                search_start = max(start + self.CHUNK_SIZE - self.CHUNK_OVERLAP, start)
                for punct in ['. ', '.\n', '! ', '? ']:
                    last_punct = text.rfind(punct, search_start, end + self.CHUNK_OVERLAP)
                    if last_punct > search_start:
                        end = last_punct + 1
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.CHUNK_OVERLAP

        return chunks

    def compare_documents(
        self,
        doc_a: ProcessedDocument,
        doc_b: ProcessedDocument,
    ) -> Dict:
        """
        Compare two documents and identify differences.

        Returns summary of comparison.
        """
        # Basic stats comparison
        comparison = {
            "doc_a": {
                "filename": doc_a.filename,
                "char_count": len(doc_a.text),
                "chunk_count": len(doc_a.chunks),
            },
            "doc_b": {
                "filename": doc_b.filename,
                "char_count": len(doc_b.text),
                "chunk_count": len(doc_b.chunks),
            },
            "size_ratio": len(doc_a.text) / len(doc_b.text) if doc_b.text else 0,
        }

        # Find common terms (simple keyword overlap)
        words_a = set(doc_a.text.lower().split())
        words_b = set(doc_b.text.lower().split())

        common_words = words_a & words_b
        only_a = words_a - words_b
        only_b = words_b - words_a

        # Filter to meaningful words (>4 chars)
        common_meaningful = {w for w in common_words if len(w) > 4}
        only_a_meaningful = {w for w in only_a if len(w) > 4}
        only_b_meaningful = {w for w in only_b if len(w) > 4}

        comparison["term_overlap"] = {
            "common_terms": len(common_meaningful),
            "only_in_a": len(only_a_meaningful),
            "only_in_b": len(only_b_meaningful),
            "overlap_ratio": len(common_meaningful) / max(len(words_a), 1),
        }

        # Sample unique terms
        comparison["sample_unique_a"] = list(only_a_meaningful)[:10]
        comparison["sample_unique_b"] = list(only_b_meaningful)[:10]

        return comparison

    def format_comparison_summary(self, comparison: Dict) -> str:
        """Format comparison results as readable text."""
        summary = f"""
## Comparacion de Documentos

### Documento A: {comparison['doc_a']['filename']}
- Caracteres: {comparison['doc_a']['char_count']:,}
- Chunks: {comparison['doc_a']['chunk_count']}

### Documento B: {comparison['doc_b']['filename']}
- Caracteres: {comparison['doc_b']['char_count']:,}
- Chunks: {comparison['doc_b']['chunk_count']}

### Analisis de Similitud
- Terminos comunes: {comparison['term_overlap']['common_terms']}
- Terminos unicos en A: {comparison['term_overlap']['only_in_a']}
- Terminos unicos en B: {comparison['term_overlap']['only_in_b']}
- Ratio de solapamiento: {comparison['term_overlap']['overlap_ratio']:.1%}

### Terminos Unicos
**Solo en Documento A:** {', '.join(comparison['sample_unique_a'][:5])}...

**Solo en Documento B:** {', '.join(comparison['sample_unique_b'][:5])}...
"""
        return summary


def main():
    """Demo of attachment handler."""
    handler = AttachmentHandler()

    print("Attachment Handler Demo")
    print(f"Supported formats: {handler.SUPPORTED_FORMATS}")
    print(f"Max file size: {handler.MAX_FILE_SIZE_MB}MB")


if __name__ == "__main__":
    main()
